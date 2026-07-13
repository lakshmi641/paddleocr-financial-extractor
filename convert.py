"""
convert.py

Stage 2 of the pipeline: turn the already-generated PaddleOCR JSON pages
(output/page_1.json ... page_N.json) into TXT, DOCX, CSV, XLSX, MD and a
merged metadata.json — WITHOUT re-running OCR.

Usage:
    python convert.py
"""

import os
import re
import glob
import json

import pandas as pd
from docx import Document
from docx.shared import Pt

OUTPUT_DIR = "output"

# Block labels, in the order they should be rendered when we fall back to
# block_order == None (headers/footers/seals rarely carry an order value).
HEADING_LABELS = {"doc_title", "paragraph_title", "figure_title"}
SKIP_LABELS = {"seal"}  # not usable as text content


def load_pages(output_dir):
    """Load every page_*.json in numeric order."""
    paths = glob.glob(os.path.join(output_dir, "page_*.json"))

    def page_num(path):
        match = re.search(r"page_(\d+)\.json$", path)
        return int(match.group(1)) if match else 0

    paths.sort(key=page_num)

    pages = []
    for path in paths:
        with open(path, "r", encoding="utf8") as f:
            pages.append((page_num(path), json.load(f)))

    return pages


def ordered_blocks(page_json):
    """Return this page's parsing_res_list sorted into reading order."""
    blocks = page_json.get("parsing_res_list", [])
    return sorted(
        blocks,
        key=lambda b: (
            b.get("block_order") if b.get("block_order") is not None else 10_000,
            b.get("block_id", 0),
        ),
    )


def html_table_to_df(html):
    try:
        dfs = pd.read_html(html)
        if not dfs:
            return None
        df = dfs[0]

        # Flatten MultiIndex / duplicate columns so downstream writers
        # (openpyxl, to_markdown) don't choke on them.
        if isinstance(df.columns, pd.MultiIndex):
            df.columns = [
                " ".join(str(c) for c in col if c not in (None, "")).strip()
                for col in df.columns
            ]

        seen = {}
        new_cols = []
        for col in df.columns:
            col = str(col)
            if col in seen:
                seen[col] += 1
                col = f"{col}_{seen[col]}"
            else:
                seen[col] = 0
            new_cols.append(col)
        df.columns = new_cols

        return df
    except Exception as e:
        print("  Table parse failed:", e)
        return None


def main():
    pages = load_pages(OUTPUT_DIR)
    if not pages:
        print(f"No page_*.json files found in '{OUTPUT_DIR}'. Nothing to convert.")
        return

    print("=" * 70)
    print(f"Converting {len(pages)} page JSON file(s) from '{OUTPUT_DIR}'")
    print("=" * 70)

    txt_lines = []
    md_lines = []
    text_rows = []          # for CSV
    tables = []              # list of (page_num, table_index, DataFrame)
    metadata_pages = []

    doc = Document()

    for page_num, page_json in pages:
        print(f"Processing page {page_num}")

        blocks = ordered_blocks(page_json)
        page_table_count = 0
        page_paragraph_count = 0

        txt_lines.append(f"\n{'=' * 60}\nPAGE {page_num}\n{'=' * 60}\n")
        md_lines.append(f"\n<!-- Page {page_num} -->\n")

        heading_added = doc.add_heading(f"Page {page_num}", level=1)
        heading_added.style.font.size = Pt(14)

        for block in blocks:
            label = block.get("block_label", "")
            content = block.get("block_content", "")

            if label in SKIP_LABELS or not content:
                continue

            if label == "table":
                df = html_table_to_df(content)
                if df is None:
                    continue

                page_table_count += 1
                tables.append((page_num, page_table_count, df))

                # DOCX table
                t = doc.add_table(rows=1, cols=len(df.columns))
                t.style = "Light Grid Accent 1"
                for i, col in enumerate(df.columns):
                    t.rows[0].cells[i].text = str(col)
                for _, row in df.iterrows():
                    cells = t.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = "" if pd.isna(val) else str(val)
                doc.add_paragraph()

                # TXT / MD table
                txt_lines.append(df.to_string(index=False) + "\n")
                md_lines.append(df.to_markdown(index=False) + "\n")
                continue

            text = " ".join(content.split()) if isinstance(content, str) else str(content)
            if not text:
                continue

            page_paragraph_count += 1
            text_rows.append({"page": page_num, "label": label, "text": text})

            if label in HEADING_LABELS:
                level = 2 if label == "doc_title" else 3
                doc.add_heading(text, level=level)
                txt_lines.append(f"\n{text}\n")
                md_lines.append(f"\n{'#' * level} {text}\n")
            else:
                doc.add_paragraph(text)
                txt_lines.append(text + "\n")
                md_lines.append(text + "\n")

        doc.add_page_break()

        metadata_pages.append(
            {
                "page": page_num,
                "width": page_json.get("width"),
                "height": page_json.get("height"),
                "paragraph_count": page_paragraph_count,
                "table_count": page_table_count,
            }
        )

    print("\nSaving outputs...")

    # DOCX
    docx_path = os.path.join(OUTPUT_DIR, "document.docx")
    doc.save(docx_path)

    # TXT
    txt_path = os.path.join(OUTPUT_DIR, "document.txt")
    with open(txt_path, "w", encoding="utf8") as f:
        f.write("\n".join(txt_lines))

    # MD
    md_path = os.path.join(OUTPUT_DIR, "document_full.md")
    with open(md_path, "w", encoding="utf8") as f:
        f.write("\n".join(md_lines))

    # CSV (paragraph text)
    csv_path = os.path.join(OUTPUT_DIR, "document.csv")
    pd.DataFrame(text_rows).to_csv(csv_path, index=False)

    # XLSX (one sheet per table, or a placeholder if none found)
    xlsx_path = os.path.join(OUTPUT_DIR, "document.xlsx")
    with pd.ExcelWriter(xlsx_path) as writer:
        if tables:
            for page_num, idx, df in tables:
                sheet_name = f"P{page_num}_T{idx}"[:31]
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        else:
            pd.DataFrame({"Info": ["No tables detected"]}).to_excel(
                writer, index=False
            )

    # metadata.json
    metadata_path = os.path.join(OUTPUT_DIR, "metadata.json")
    with open(metadata_path, "w", encoding="utf8") as f:
        json.dump(
            {
                "page_count": len(pages),
                "paragraph_count": len(text_rows),
                "table_count": len(tables),
                "pages": metadata_pages,
            },
            f,
            indent=4,
            ensure_ascii=False,
        )

    print()
    print("=" * 70)
    print("Completed")
    print("=" * 70)
    print(f"Pages      : {len(pages)}")
    print(f"Paragraphs : {len(text_rows)}")
    print(f"Tables     : {len(tables)}")
    print(f"DOCX       : {docx_path}")
    print(f"TXT        : {txt_path}")
    print(f"CSV        : {csv_path}")
    print(f"XLSX       : {xlsx_path}")
    print(f"MD         : {md_path}")
    print(f"Metadata   : {metadata_path}")


if __name__ == "__main__":
    main()
